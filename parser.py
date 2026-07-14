import re
import io
import pdfplumber
from docx import Document

def extract_text_from_image(file_input):
    """
    Extract text from an image file (JPG, PNG, WEBP, BMP, TIFF) using pytesseract OCR.
    Accepts file path string, raw bytes, or file-like object.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ValueError(
            "Image OCR requires pytesseract and Pillow. "
            "Install with: pip install pytesseract Pillow  "
            "and ensure Tesseract is installed on your system."
        )

    if isinstance(file_input, str):
        img = Image.open(file_input)
    elif isinstance(file_input, bytes):
        img = Image.open(io.BytesIO(file_input))
    else:
        data = file_input.read() if hasattr(file_input, 'read') else file_input.getvalue()
        img = Image.open(io.BytesIO(data))

    try:
        text = pytesseract.image_to_string(img)
    except Exception:
        raise ValueError(
            "Could not parse the image. If this is a generic image (like a photo or screenshot), "
            "please upload a proper document (PDF or DOCX) containing your resume instead."
        )
    return text.strip()

def extract_text_from_pdf(file_input):
    """
    Extract raw text from a PDF file.
    Accepts file path or file-like object.
    """
    text = ""
    # Check if input is a string path or file-like object
    if isinstance(file_input, str):
        with pdfplumber.open(file_input) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    else:
        # File-like object (e.g., Streamlit UploadedFile)
        with pdfplumber.open(file_input) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(file_input):
    """
    Extract raw text from a DOCX file.
    Accepts file path or file-like object.
    """
    if isinstance(file_input, str):
        doc = Document(file_input)
    else:
        # Wrap bytes in a BytesIO if it's not already file-like or if python-docx requires it
        if isinstance(file_input, bytes):
            doc = Document(io.BytesIO(file_input))
        else:
            doc = Document(file_input)
            
    # Extract text from paragraphs
    paragraphs_text = [para.text for para in doc.paragraphs]
    
    # Also extract text from tables, which is often neglected
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            table_text.append(" | ".join(row_text))
            
    full_text = "\n".join(paragraphs_text + table_text)
    return full_text.strip()

def segment_resume(text):
    """
    Segment resume text into logical sections based on common headings.
    Returns a dictionary of section names to lists of string lines/chunks.
    Sections: skills, experience, education, projects, summary/other.
    """
    sections = {
        "summary": [],
        "skills": [],
        "experience": [],
        "education": [],
        "projects": []
    }
    
    # Regex patterns for common headers
    headers_patterns = {
        "skills": re.compile(r'^(technical\s+)?skills(?!\s+and\s+experience)', re.IGNORECASE),
        "experience": re.compile(r'^(work\s+)?experience|employment|work\s+history|career\s+history|professional\s+experience', re.IGNORECASE),
        "education": re.compile(r'^education|academic\s+background|credentials', re.IGNORECASE),
        "projects": re.compile(r'^projects|personal\s+projects|academic\s+projects|open\s+source', re.IGNORECASE)
    }
    
    lines = [line.strip() for line in text.split('\n')]
    current_section = "summary" # Default starting section is summary/intro
    
    for line in lines:
        if not line:
            continue
            
        # Check if the line is a section header
        # Headers are usually short and match our regex patterns
        is_header = False
        if len(line) < 40:
            for sec_name, pattern in headers_patterns.items():
                if pattern.match(line.strip(' :-•*')):
                    current_section = sec_name
                    is_header = True
                    break
        
        if is_header:
            continue
            
        sections[current_section].append(line)
        
    cleaned_sections = {}
    for key, val in sections.items():
        text_block = "\n".join(val)
        
        # If it is the skills section, we split by commas, newlines, or pipes
        if key == "skills":
            chunks = []
            for chunk in re.split(r'[\n\|,]+', text_block):
                chunk = chunk.strip().lstrip('•-*+o ')
                if len(chunk) >= 2:  # Allow shorter words like "Go", "R"
                    chunks.append(chunk)
        else:
            # For other sections, split by newlines
            chunks = []
            for chunk in re.split(r'\n+', text_block):
                chunk = chunk.strip().lstrip('•-*+o ')
                if len(chunk) > 10:  # Ignore very short fragments
                    chunks.append(chunk)
                    
        cleaned_sections[key] = chunks
        
    return cleaned_sections

def parse_job_description(text):
    """
    Parse a JD plain text into required skills, preferred skills, and responsibilities.
    Returns a dictionary.
    """
    sections = {
        "required_skills": [],
        "preferred_skills": [],
        "responsibilities": [],
        "general": []
    }
    
    # Split JD into sections
    headers_patterns = {
        "required_skills": re.compile(r'required\s+(skills|qualifications|experience)|requirements|what\s+you\s+need|must\s+have', re.IGNORECASE),
        "preferred_skills": re.compile(r'preferred\s+(skills|qualifications|experience)|nice\s+to\s+have|plus|desired|preferred', re.IGNORECASE),
        "responsibilities": re.compile(r'responsibilities|role|duties|what\s+you\s+will\s+do|key\s+responsibilities', re.IGNORECASE)
    }
    
    lines = [line.strip() for line in text.split('\n')]
    current_section = "general"
    
    for line in lines:
        if not line:
            continue
            
        is_header = False
        # If line is short and doesn't start with bullet, check if it's a header
        if len(line) < 45 and not line.strip().startswith(('-', '•', '*', '+')):
            matched_target = False
            for sec_name, pattern in headers_patterns.items():
                if pattern.search(line.strip(' :-•*')):
                    current_section = sec_name
                    is_header = True
                    matched_target = True
                    break
            
            # If it's a short line/header but doesn't match targets, reset to general (e.g. Benefits, About Us)
            if not matched_target and (line.isupper() or line.endswith(':') or len(line) < 25):
                current_section = "general"
                is_header = True
                    
        if is_header:
            continue
            
        sections[current_section].append(line)
        
    # Process each section into individual requirement chunks (bullet points/sentences)
    cleaned_sections = {}
    for key, val in sections.items():
        text_block = "\n".join(val)
        chunks = []
        for chunk in re.split(r'\n+', text_block):
            chunk = chunk.strip().lstrip('•-*+o ')
            if len(chunk) > 10:
                chunks.append(chunk)
        cleaned_sections[key] = chunks
        
    return cleaned_sections
